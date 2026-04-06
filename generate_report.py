#!/usr/bin/env python3
"""Generate Project Report as DOCX"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Title
title = doc.add_heading('ML-Based CI/CD Failure Prediction System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Project Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'\nGenerated on: {datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph().add_run().add_break()

# Executive Summary
doc.add_heading('1. Executive Summary', 1)
doc.add_paragraph(
    'This project implements a production-ready Machine Learning service that predicts '
    'Continuous Integration/Continuous Deployment (CI/CD) build outcomes before full '
    'pipeline execution. By analyzing commit-level metadata, the system provides early '
    'risk assessment, enabling development teams to reduce wasted computational resources, '
    'accelerate feedback loops, and improve overall pipeline efficiency.'
)

doc.add_paragraph(
    'The solution features a cloud-native architecture with multiple deployment options, '
    'comprehensive observability through Prometheus and Grafana, and seamless integration '
    'with popular CI/CD platforms like GitHub Actions, GitLab CI, and Jenkins.'
)

# Problem Statement
doc.add_heading('2. Problem Statement', 1)
doc.add_paragraph(
    'In modern software development, CI/CD pipelines often fail after spending significant '
    'time on installation, testing, and build phases. Common issues include:'
)

p = doc.add_paragraph(style='List Bullet')
p.add_run('Late failure detection consuming valuable compute resources')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Extended feedback cycles delaying developer productivity')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Unpredictable pipeline behavior causing deployment delays')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Inefficient resource utilization in cloud environments')

doc.add_paragraph(
    'This project addresses these challenges by introducing an intelligent prediction '
    'layer that estimates build success probability before resource-intensive operations.'
)

# Solution Overview
doc.add_heading('3. Solution Overview', 1)

doc.add_heading('3.1 Core Components', 2)
components = [
    ('Flask REST API', 'Prediction endpoints with caching and metrics'),
    ('Random Forest Model', 'Pre-trained ML model with 85% accuracy'),
    ('Redis Cache', 'Prediction result caching for performance'),
    ('S3 Storage', 'Cloud-based model artifact management'),
    ('Lambda Function', 'Serverless prediction option'),
    ('Monitoring Stack', 'Prometheus metrics and Grafana dashboards')
]

for name, desc in components:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

# Architecture
doc.add_heading('3.2 System Architecture', 2)
doc.add_paragraph(
    'The system follows a microservices architecture with the following flow:\n\n'
    '1. CI/CD platform triggers prediction request with commit metadata\n'
    '2. API checks Redis cache for existing predictions\n'
    '3. On cache miss, ML model processes the request\n'
    '4. Result is cached and returned with risk assessment\n'
    '5. Metrics are collected for monitoring and analysis'
)

# Technical Specifications
doc.add_heading('4. Technical Specifications', 1)

doc.add_heading('4.1 Model Details', 2)
model_table = doc.add_table(rows=7, cols=2)
model_table.style = 'Light Grid Accent 1'

model_data = [
    ['Attribute', 'Value'],
    ['Algorithm', 'Random Forest Classifier'],
    ['Training Dataset', '2,000 CI/CD build samples'],
    ['Input Features', '8 (commit size, files changed, test coverage, etc.)'],
    ['Output Classes', 'PASS / FAIL'],
    ['Test Accuracy', '~85%'],
    ['Model Size', '5.4 MB (serialized)']
]

for i, (key, value) in enumerate(model_data):
    row_cells = model_table.rows[i].cells
    row_cells[0].text = key
    row_cells[1].text = value
    if i == 0:
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_heading('4.2 Input Features', 2)
doc.add_paragraph('The model analyzes the following commit-level features:')

features = [
    ('commit_size', 'Number of lines changed in commit'),
    ('files_changed', 'Count of modified files'),
    ('test_coverage', 'Percentage of code covered by tests'),
    ('past_failures', 'Number of recent build failures'),
    ('dependency_changes', 'Count of dependency file modifications'),
    ('author_experience', 'Developer\'s commit history count'),
    ('time_of_commit', 'Hour of day when commit was made'),
    ('build_time', 'Historical average build duration')
]

for feature, description in features:
    p = doc.add_paragraph()
    p.add_run(f'{feature}: ').bold = True
    p.add_run(description)

# Deployment Options
doc.add_heading('5. Deployment Options', 1)

deployments = [
    ('Docker Compose',
     'Local development with full stack (API, Redis, Prometheus, Grafana). '
     'Recommended for testing and demonstrations.'),

    ('Kubernetes',
     'Production-ready container orchestration with deployment manifests '
     'and Helm chart for cloud-native deployments.'),

    ('AWS Infrastructure (Terraform)',
     'Complete AWS setup including ECS, Lambda, S3, VPC, and CloudWatch. '
     'Includes infrastructure as code for reproducible deployments.'),

    ('Serverless (Lambda)',
     'Event-driven prediction function with API Gateway. '
     'Pay-per-use pricing with automatic scaling.')
]

for name, desc in deployments:
    doc.add_heading(name, 2)
    doc.add_paragraph(desc)

# API Documentation
doc.add_heading('6. API Documentation', 1)

doc.add_heading('6.1 Endpoints', 2)
api_table = doc.add_table(rows=5, cols=3)
api_table.style = 'Light Grid Accent 1'

api_data = [
    ['Method', 'Endpoint', 'Description'],
    ['GET', '/health', 'Service health check'],
    ['POST', '/predict', 'Single build prediction'],
    ['POST', '/predict-batch', 'Batch predictions for multiple commits'],
    ['GET', '/metrics', 'Prometheus metrics endpoint']
]

for i, row_data in enumerate(api_data):
    row = api_table.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# Monitoring & Observability
doc.add_heading('7. Monitoring & Observability', 1)

doc.add_heading('7.1 Metrics Collection', 2)
metrics = [
    'prediction_requests_total - Total prediction requests',
    'prediction_duration_seconds - Request processing time',
    'cache_hits_total - Redis cache hit count',
    'cache_misses_total - Redis cache miss count',
    'model_load_time_seconds - Model loading duration'
]

for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

doc.add_heading('7.2 Dashboards', 2)
doc.add_paragraph(
    'Grafana provides real-time visualization of:\n'
    '• Prediction success rate and failure distribution\n'
    '• API response time percentiles\n'
    '• Cache hit/miss ratios\n'
    '• Resource utilization (CPU, memory)'
)

# CI/CD Integration
doc.add_heading('8. CI/CD Integration', 1)

doc.add_paragraph(
    'The service integrates seamlessly with popular CI/CD platforms:'
)

integrations = [
    ('GitHub Actions',
     'Workflow file provided in ci_cd_workflow.yml. '
     'Stops pipeline early if high failure risk detected.'),

    ('GitLab CI',
     'HTTP request to prediction API with conditional job execution.'),

    ('Jenkins',
     'Pipeline stage with HTTP request and conditional logic.'),

    ('CircleCI',
     'Orb integration for prediction-based workflow control.')
]

for platform, desc in integrations:
    p = doc.add_paragraph()
    p.add_run(f'{platform}: ').bold = True
    p.add_run(desc)

# Project Structure
doc.add_heading('9. Project Structure', 1)

doc.add_paragraph('Key directories and files:')
structure = '''
├── app.py                    # Flask API with prediction endpoints
├── models/                   # ML model artifacts (.pkl files)
├── k8s/                      # Kubernetes manifests
├── helm/                     # Helm chart for deployment
├── terraform/                # AWS infrastructure code
├── lambda/                   # AWS Lambda function
├── docker-compose.yml        # Multi-service orchestration
├── s3_model_loader.py        # S3 model loading logic
├── redis_cache.py            # Redis caching layer
├── prometheus_metrics.py     # Metrics collection
└── grafana/                  # Dashboard configurations
'''

doc.add_paragraph(structure)

# Technology Stack
doc.add_heading('10. Technology Stack', 1)

tech_table = doc.add_table(rows=7, cols=2)
tech_table.style = 'Light Grid Accent 1'

tech_data = [
    ['Category', 'Technologies'],
    ['Backend', 'Flask, Python 3.9+'],
    ['ML/AI', 'scikit-learn, pandas, numpy'],
    ['Caching', 'Redis 7'],
    ['Monitoring', 'Prometheus, Grafana'],
    ['Containers', 'Docker, Kubernetes, Helm'],
    ['Cloud', 'AWS (ECS, Lambda, S3, CloudWatch), Terraform']
]

for i, (cat, tech) in enumerate(tech_data):
    row_cells = tech_table.rows[i].cells
    row_cells[0].text = cat
    row_cells[1].text = tech
    if i == 0:
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# Future Enhancements
doc.add_heading('11. Future Enhancements', 1)
enhancements = [
    'Support for additional ML algorithms (XGBoost, Neural Networks)',
    'A/B testing framework for model comparison',
    'Real-time model retraining pipeline',
    'Multi-language SDK support (JavaScript, Go)',
    'Extended feature set with code complexity metrics',
    'Integration with SonarQube for code quality analysis'
]

for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

# Conclusion
doc.add_heading('12. Conclusion', 1)
doc.add_paragraph(
    'This project demonstrates a practical application of Machine Learning in DevOps '
    'workflows. By providing early failure prediction, it helps teams optimize their '
    'CI/CD pipelines, reduce costs, and improve developer experience. The cloud-native '
    'architecture ensures scalability and flexibility for various deployment scenarios.'
)

# References
doc.add_heading('13. References', 1)
references = [
    'scikit-learn: https://scikit-learn.org/',
    'Flask: https://flask.palletsprojects.com/',
    'Docker: https://docs.docker.com/',
    'Kubernetes: https://kubernetes.io/docs/',
    'Terraform: https://www.terraform.io/docs',
    'Prometheus: https://prometheus.io/docs/',
    'GitHub Repository: https://github.com/practicalClerk/Emerging_Trends_Tools'
]

for ref in references:
    doc.add_paragraph(ref, style='List Bullet')

# Save document
output_path = '/Users/dhruvvaishnav/Downloads/Devops_Project-main/PROJECT_REPORT.docx'
doc.save(output_path)
print(f"✅ Project report generated: {output_path}")
